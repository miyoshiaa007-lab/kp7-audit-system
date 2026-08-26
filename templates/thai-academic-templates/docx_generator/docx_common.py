# -*- coding: utf-8 -*-
"""
docx_common.py
---------------
โมดูลกลางสำหรับสร้างเทมเพลตเอกสารวิชาการไทย (งานวิจัย 5 บท และ ผลงานทางวิชาการ/PA)
ด้วย python-docx ให้ตรงตามสเปกที่สรุปไว้ใน analysis/STRUCTURE_ANALYSIS.md

ครอบคลุม:
- การตั้งค่าฟอนต์เริ่มต้นทั้งเอกสาร (TH Sarabun New)
- สไตล์หัวข้อ (บทที่ / หัวข้อรอง / หัวข้อย่อย)
- การตั้งค่าหน้ากระดาษ/ระยะขอบ ตามตัวอย่างที่แนบ
- ระบบเลขหน้า 2 รูปแบบในเล่มเดียว: ส่วนหน้า = พยัญชนะไทย (ก ข ค…),
  ส่วนเนื้อหา = เลขอารบิก (1 2 3…) โดยใช้ OOXML `w:pgNumType/@w:fmt`
  ค่า "thaiLetters" / "decimal" ซึ่งเป็นฟีเจอร์มาตรฐานของ Word เอง
  (ไม่ใช่การพิมพ์ตัวอักษรปลอมด้วยมือ)
- ตารางแบบเส้นเต็มหัวตารางสีเทาอ่อน ตรงตามตัวอย่าง PA 4 ที่แนบ
- กล่อง placeholder สำหรับรูปภาพ/หลักฐาน และบล็อกลงนาม
"""
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from thai_text_utils import clean_for_docx

FONT_TH = "TH Sarabun New"

# ---- ขนาดฟอนต์มาตรฐาน (pt) ----
SIZE_BODY = 16
SIZE_H3 = 16
SIZE_H2 = 18
SIZE_H1 = 20
SIZE_SMALL_NOTE = 14

# ---- ระยะขอบกระดาษ (ซม.) ตามไฟล์ตัวอย่างที่แนบ ----
MARGIN_TOP = 3.0
MARGIN_BOTTOM = 2.54
MARGIN_LEFT = 3.81   # เผื่อเข้าเล่ม
MARGIN_RIGHT = 2.54
HEADER_DIST = 1.25
FOOTER_DIST = 1.25

# พยัญชนะไทยที่ใช้เรียงลำดับเลขหน้าส่วนหน้า (ก ข ค ...) - 44 ตัวตามมาตรฐาน OOXML thaiLetters
THAI_LETTERS = list("กขคงจฉชซฌญฎฏฐฑฒณดตถทธนบปผฝพฟภมยรลวศษสหฬอฮ")


# =========================================================================
# ฟอนต์ / ภาษา
# =========================================================================

def _set_run_thai_font(run, size=SIZE_BODY, bold=False, color=None):
    run.font.name = FONT_TH
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(attr), FONT_TH)
    lang = rPr.find(qn('w:lang'))
    if lang is None:
        lang = OxmlElement('w:lang')
        rPr.append(lang)
    lang.set(qn('w:val'), 'th-TH')
    lang.set(qn('w:eastAsia'), 'th-TH')
    lang.set(qn('w:bidi'), 'th-TH')


def set_document_defaults(doc: Document):
    """ตั้งฟอนต์เริ่มต้น (Normal) ของทั้งเอกสารเป็น TH Sarabun New 16pt"""
    style = doc.styles['Normal']
    style.font.name = FONT_TH
    style.font.size = Pt(SIZE_BODY)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(attr), FONT_TH)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)


# =========================================================================
# ย่อหน้า / หัวข้อ
# =========================================================================

def add_paragraph_thai(doc_or_cell, text="", size=SIZE_BODY, bold=False,
                        align=WD_ALIGN_PARAGRAPH.THAI_JUSTIFY, indent_first=None,
                        space_after=6, color=None, style=None):
    """เพิ่มย่อหน้าข้อความไทย พร้อมแทรก ZWS และตั้งค่าฟอนต์/ภาษาไทยให้ทุก run"""
    p = doc_or_cell.add_paragraph(style=style)
    p.alignment = align
    if indent_first is not None:
        p.paragraph_format.first_line_indent = Cm(indent_first)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(clean_for_docx(text))
        _set_run_thai_font(run, size=size, bold=bold, color=color)
    return p


def add_mixed_paragraph(doc_or_cell, parts, align=WD_ALIGN_PARAGRAPH.THAI_JUSTIFY,
                         space_after=6, indent_first=None):
    """
    parts: list of (text, bold) เพื่อสร้างย่อหน้าที่มีบางส่วนตัวหนา
    เช่น [("ผลการดำเนินงาน : ", True), ("รายละเอียด...", False)]
    """
    p = doc_or_cell.add_paragraph()
    p.alignment = align
    if indent_first is not None:
        p.paragraph_format.first_line_indent = Cm(indent_first)
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold in parts:
        run = p.add_run(clean_for_docx(text))
        _set_run_thai_font(run, size=SIZE_BODY, bold=bold)
    return p


def add_heading1(doc, text, number=None):
    """หัวข้อระดับบทที่ (20pt ตัวหนา กึ่งกลาง) - ใช้ page break ก่อนหน้าเสมอ"""
    doc.add_page_break()
    if number:
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(number)
        _set_run_thai_font(r1, size=SIZE_H1, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(clean_for_docx(text))
    _set_run_thai_font(r2, size=SIZE_H1, bold=True)
    p2.paragraph_format.space_after = Pt(18)
    return p2


def add_heading2(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(clean_for_docx(text))
    _set_run_thai_font(r, size=SIZE_H2, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_heading3(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(clean_for_docx(text))
    _set_run_thai_font(r, size=SIZE_H3, bold=True)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_photo_placeholder(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(clean_for_docx(
        "(พื้นที่สำหรับภาพความสำเร็จ / ร่องรอยหลักฐานเชิงประจักษ์ "
        "— แทรกภาพถ่ายกิจกรรม เอกสาร หรือรหัสคิวอาร์ลิงก์หลักฐานที่นี่)"))
    _set_run_thai_font(r, size=SIZE_SMALL_NOTE, bold=False, color=(0x59, 0x59, 0x59))
    p.paragraph_format.space_after = Pt(12)
    # กรอบสี่เหลี่ยมรอบย่อหน้า (ให้เห็นเป็นกล่องชัดเจนสำหรับผู้กรอกแบบ)
    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement('w:pBdr')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), '999999')
        borders.append(el)
    # w:pBdr ต้องอยู่ก่อน w:jc/w:spacing ฯลฯ ตามลำดับ schema ของ CT_PPr
    pbdr_successors = ('w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku',
                        'w:wordWrap', 'w:overflowPunct', 'w:topLinePunct',
                        'w:autoSpaceDE', 'w:autoSpaceDN', 'w:bidi', 'w:adjustRightInd',
                        'w:snapToGrid', 'w:spacing', 'w:ind', 'w:contextualSpacing',
                        'w:mirrorIndents', 'w:suppressOverlap', 'w:jc', 'w:textDirection',
                        'w:textAlignment', 'w:textboxTightWrap', 'w:outlineLvl', 'w:divId',
                        'w:cnfStyle', 'w:rPr', 'w:sectPr', 'w:pPrChange')
    pPr.insert_element_before(borders, *pbdr_successors)
    return p


def add_disclaimer_note(doc, text=None):
    text = text or ("*** ตัวเลขทั้งหมดในตารางนี้เป็นข้อมูลสมมุติ (ตัวอย่าง) "
                     "เพื่อแสดงวิธีการกรอกและการคำนวณเท่านั้น — ต้องแทนที่ด้วยข้อมูลจริงก่อนยื่นเอกสาร ***")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(clean_for_docx(text))
    _set_run_thai_font(r, size=SIZE_SMALL_NOTE, bold=True, color=(0xC0, 0x00, 0x00))
    p.paragraph_format.space_after = Pt(6)
    return p


def add_signature_block(doc, name="( ......................................... )",
                         position="ตำแหน่ง ................ วิทยฐานะ ................",
                         org="................................................"):
    doc.add_paragraph()
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(clean_for_docx(
        "ลงชื่อ . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . ผู้รายงาน"))
    _set_run_thai_font(r1, size=SIZE_BODY)
    for text in (name, position, org):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(clean_for_docx(text))
        _set_run_thai_font(r, size=SIZE_BODY)


# =========================================================================
# ตาราง
# =========================================================================

def _shade_cell(cell, hex_color="D9D9D9"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tcPr.append(borders)


def _set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=SIZE_BODY):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    if text:
        run = p.add_run(clean_for_docx(str(text)))
        _set_run_thai_font(run, size=size, bold=bold)
    _set_cell_borders(cell)


def add_table_grid(doc, headers, rows, col_widths_cm=None,
                    header_align=WD_ALIGN_PARAGRAPH.CENTER,
                    body_align=WD_ALIGN_PARAGRAPH.LEFT,
                    center_cols=None):
    """
    สร้างตารางเส้นเต็ม หัวตารางตัวหนา พื้นเทาอ่อน กึ่งกลาง
    headers: list[str]
    rows: list[list[str]]
    center_cols: index คอลัมน์ (0-based) ที่ต้องการจัดกึ่งกลางในส่วนเนื้อหา (เช่น คอลัมน์ตัวเลข/ลำดับ)
    """
    center_cols = center_cols or []
    n_cols = len(headers)
    table = doc.add_table(rows=1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_text(hdr_cells[i], h, bold=True, align=header_align)
        _shade_cell(hdr_cells[i])

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if i in center_cols else body_align
            _set_cell_text(cells[i], val, bold=False, align=align)

    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for r in table.rows:
                r.cells[i].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


# =========================================================================
# หน้ากระดาษ / เลขหน้า
# =========================================================================

def set_page_setup(section, top=MARGIN_TOP, bottom=MARGIN_BOTTOM,
                    left=MARGIN_LEFT, right=MARGIN_RIGHT,
                    header=HEADER_DIST, footer=FOOTER_DIST):
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)
    section.header_distance = Cm(header)
    section.footer_distance = Cm(footer)


def set_page_number_format(section, fmt="decimal", start=None):
    """
    ตั้งรูปแบบเลขหน้าของ section (มาตรฐาน OOXML w:pgNumType):
      fmt="thaiLetters" -> ก ข ค ง ...
      fmt="decimal"      -> 1 2 3 ...
    """
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        # ตำแหน่งที่ถูกต้องตามลำดับ schema ของ CT_SectPr คือ "ก่อน" w:cols
        # (และองค์ประกอบถัดจากนี้) มิฉะนั้น Word จะถือว่าไฟล์เสียหายและเสนอซ่อมไฟล์
        successors = ('w:cols', 'w:formProt', 'w:vAlign', 'w:noEndnote', 'w:titlePg',
                      'w:textDirection', 'w:bidi', 'w:rtlGutter', 'w:docGrid',
                      'w:printerSettings', 'w:sectPrChange')
        sectPr.insert_element_before(pgNumType, *successors)
    pgNumType.set(qn('w:fmt'), fmt)
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))


def suppress_first_page_footer(section):
    """ซ่อนเลขหน้าบนหน้าแรกของ section (ใช้กับหน้าปกที่ไม่ต้องการแสดงเลขหน้า)"""
    section.different_first_page_header_footer = True
    # ล้างเนื้อหา footer ของหน้าแรกให้ว่าง (ไม่ต้องแสดงอะไร)
    ffooter = section.first_page_footer
    for p in list(ffooter.paragraphs):
        p.clear()


def add_page_number_footer(section, size=SIZE_SMALL_NOTE):
    """ใส่ field PAGE กึ่งกลางท้ายกระดาษ (รูปแบบจะเปลี่ยนไปตาม pgNumType ของ section นั้น ๆ โดยอัตโนมัติ)"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    _set_run_thai_font(run, size=size)
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    cached_text = OxmlElement('w:t')
    cached_text.text = '1'
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_sep)
    run._r.append(cached_text)
    run._r.append(fldChar_end)
    return footer


def new_section(doc, start_type=WD_SECTION.NEW_PAGE):
    section = doc.add_section(start_type)
    set_page_setup(section)
    return section


# =========================================================================
# สารบัญอัตโนมัติ (TOC field) - ผู้ใช้กด F9 / Update Field ใน Word เพื่ออัปเดต
# =========================================================================

def add_toc_field(doc, caption="สารบัญ"):
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(clean_for_docx(caption))
    _set_run_thai_font(r_title, size=SIZE_H1, bold=True)
    p_title.paragraph_format.space_after = Pt(18)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    _set_run_thai_font(run, size=SIZE_BODY)
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    fldChar_begin.set(qn('w:dirty'), 'true')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')

    note_run_text = OxmlElement('w:t')
    note_run_text.text = "คลิกขวา > Update Field เพื่ออัปเดตสารบัญหลังพิมพ์เนื้อหาเสร็จ"

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar_begin)
    r_element.append(instrText)
    r_element.append(fldChar_sep)
    r_element.append(note_run_text)
    r_element.append(fldChar_end)
    return paragraph
