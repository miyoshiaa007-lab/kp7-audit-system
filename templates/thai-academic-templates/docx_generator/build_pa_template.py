# -*- coding: utf-8 -*-
"""
build_pa_template.py
----------------------
สร้างเทมเพลต Word (.docx) สำหรับ "ผลงานทางวิชาการ / รายงานผลการปฏิบัติงานตามข้อตกลง
ในการพัฒนางาน (PA)" ตามโครงสร้างที่พบในไฟล์ตัวอย่างที่แนบ (PA4 ส่วนที่ 2 และส่วนที่ 3
ของศึกษานิเทศก์) โดยแปลงเป็น "เทมเพลตเปล่า" ที่มีโครงสร้าง/หัวตาราง/จำนวนแถวตรงตาม
ตัวอย่าง ให้ผู้ใช้กรอกข้อมูลจริงของตนเองแทนที่ข้อความในวงเล็บเหลี่ยม [ ... ]

โครงสร้างที่ทำซ้ำได้ต่อ 1 ตัวชี้วัด (ประกอบเป็นฟังก์ชัน build_indicator_block):
  1. หัวเรื่องส่วน/ด้าน/ตัวชี้วัด
  2. ย่อหน้าแนะนำตัว + ภาพรวมกลยุทธ์/รูปแบบ (bullet)
  3. เกณฑ์การให้คะแนน (Scoring Rubric) 1-5 คะแนน
  4. ตารางสรุปความสอดคล้องระหว่างเกณฑ์ PA 4 กับกลยุทธ์
  5. การดำเนินงานตามกลยุทธ์ (ขั้นเตรียมการ/ขั้นดำเนินการ PDCA/ขั้นจัดทำรายงาน)
  6. ตารางที่ 1 วิเคราะห์ความสอดคล้องกับเกณฑ์
  7. ผลลัพธ์ + ตารางที่ 2 (มีคำเตือนตัวเลขสมมุติ)
  8. รายการหลักฐานเชิงประจักษ์รายข้อ (x.x.1 ... ) พร้อมกล่องแทรกภาพ
  9. ตารางที่ 3 สรุปเทียบค่าเป้าหมาย (มีคำเตือนตัวเลขสมมุติ)
  10. ตารางแหล่งข้อมูลสำหรับตรวจสอบ/แทนที่ตัวเลขสมมุติ (ลบก่อนยื่นจริง)
  11. บล็อกลงชื่อผู้รายงาน

รัน:  python build_pa_template.py
ผลลัพธ์: ../output/PA-Academic-Report-Template.docx
"""
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

from docx_common import (
    set_document_defaults, set_page_setup, set_page_number_format,
    add_page_number_footer, add_heading2, add_heading3, add_paragraph_thai,
    add_mixed_paragraph, add_table_grid, add_photo_placeholder,
    add_disclaimer_note, add_signature_block, add_toc_field, new_section,
    THAI_LETTERS, SIZE_H1, SIZE_BODY, _set_run_thai_font,
)

OUT_PATH = os.path.join(os.path.dirname(__file__), "..", "output",
                         "PA-Academic-Report-Template.docx")


def _center_run(doc, text, size=SIZE_BODY, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    _set_run_thai_font(r, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    return p


# ---------------------------------------------------------------------------
# หน้าปก / ส่วนที่ 1 ข้อมูลทั่วไป
# ---------------------------------------------------------------------------

def build_cover_page(doc):
    for _ in range(3):
        doc.add_paragraph()
    _center_run(doc, "รายงานผลการปฏิบัติงานตามข้อตกลงในการพัฒนางาน (PA)", size=22, bold=True, space_after=4)
    _center_run(doc, "(วฐ. 2 / PA 4)", size=18, bold=True, space_after=40)
    _center_run(doc, "ของ", size=SIZE_H1, space_after=16)
    _center_run(doc, "[คำนำหน้า ชื่อ-นามสกุลผู้จัดทำ]", size=SIZE_H1, bold=True, space_after=6)
    _center_run(doc, "ตำแหน่ง [ตำแหน่ง] วิทยฐานะ [วิทยฐานะที่ขอรับการประเมิน]", size=SIZE_BODY, space_after=40)
    _center_run(doc, "[ชื่อสถานศึกษา/หน่วยงาน]", size=SIZE_BODY)
    _center_run(doc, "[ชื่อสำนักงานเขตพื้นที่การศึกษา/สังกัด]", size=SIZE_BODY)
    _center_run(doc, "รอบการประเมิน [วันที่ 1 ตุลาคม พ.ศ. .... – 30 กันยายน พ.ศ. ....]", size=SIZE_BODY, space_after=0)


def build_part1_general_info(doc):
    _center_run(doc, "ส่วนที่ 1", size=SIZE_H1, bold=True, space_after=2)
    _center_run(doc, "ข้อมูลทั่วไปและข้อตกลงในการพัฒนางาน", size=SIZE_H1, bold=True, space_after=18)
    for label, value in [
        ("ชื่อ-นามสกุล", "[คำนำหน้า ชื่อ-นามสกุล]"),
        ("ตำแหน่ง", "[ตำแหน่ง]"),
        ("วิทยฐานะปัจจุบัน", "[วิทยฐานะปัจจุบัน]"),
        ("วิทยฐานะที่ขอรับการประเมิน", "[วิทยฐานะที่ขอรับการประเมิน]"),
        ("สังกัด", "[หน่วยงาน/สถานศึกษา/สำนักงานเขตพื้นที่การศึกษา]"),
        ("รอบการประเมิน", "[วันที่เริ่มต้น – วันที่สิ้นสุด]"),
    ]:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}  ")
        _set_run_thai_font(r1, size=SIZE_BODY, bold=True)
        r2 = p.add_run(value)
        _set_run_thai_font(r2, size=SIZE_BODY)
    doc.add_paragraph()
    add_paragraph_thai(doc, "(สรุปข้อตกลงในการพัฒนางาน (PA 1) ที่ได้จัดทำไว้ล่วงหน้ากับผู้บังคับบัญชา "
                             "โดยย่อ ระบุประเด็นท้าทายและตัวชี้วัดความสำเร็จที่ตกลงไว้)",
                        indent_first=1.25)


# ---------------------------------------------------------------------------
# บล็อกตัวชี้วัด (ทำซ้ำได้ทุกด้าน/ทุกตัวชี้วัด)
# ---------------------------------------------------------------------------

def build_indicator_block(doc, part_label, domain_title, indicator_no, indicator_title,
                           num_criteria=5, num_strategies=8, num_source_rows=5,
                           leading_pagebreak=True):
    """
    part_label: เช่น "ส่วนที่ 2"
    domain_title: เช่น "ผลการดำเนินงาน ด้านที่ 1 ด้านทักษะการวางแผนพัฒนาการนิเทศการศึกษา ..."
    indicator_no, indicator_title: เช่น 1, "กระบวนการจัดการเรียนรู้ (ประกันคุณภาพ / การพัฒนาคุณภาพอย่างต่อเนื่อง)"
    leading_pagebreak: ใส่ False เมื่อบล็อกนี้เป็นเนื้อหาแรกสุดของไฟล์ (เช่น ไฟล์แยกต่อ
        ตัวชี้วัด) เพื่อไม่ให้มีหน้าว่างเปล่าแทรกอยู่หน้าแรก
    """
    if leading_pagebreak:
        doc.add_page_break()
    _center_run(doc, part_label, size=SIZE_H1, bold=True, space_after=2)
    _center_run(doc, domain_title, size=18, bold=True, space_after=2)
    _center_run(doc, f"ตัวชี้วัดที่ {indicator_no}  {indicator_title}", size=18, bold=True, space_after=16)

    # 1) ย่อหน้าแนะนำตัว
    add_paragraph_thai(doc,
        "ข้าพเจ้า [คำนำหน้า ชื่อ-นามสกุล] ตำแหน่ง [ตำแหน่ง] วิทยฐานะ [วิทยฐานะ] "
        f"สังกัด [หน่วยงาน] ขอรายงานผลการดำเนินงาน{domain_title} "
        f"ตัวชี้วัดที่ {indicator_no} {indicator_title} ดังรายละเอียดต่อไปนี้")

    # 2) ภาพรวมรูปแบบ/กลยุทธ์ (bullet placeholder)
    add_paragraph_thai(doc, "ข้าพเจ้าดำเนินงานภายใต้รูปแบบ [ชื่อรูปแบบ/นวัตกรรมของผู้จัดทำ] "
                             f"ซึ่งประกอบด้วย {num_strategies} กลยุทธ์ ดังนี้")
    for i in range(1, num_strategies + 1):
        add_paragraph_thai(doc, f"•  กลยุทธ์ที่ {i}  [ชื่อกลยุทธ์ที่ {i}]  [คำอธิบายสั้น ๆ]", indent_first=None)

    # 3) เกณฑ์การให้คะแนน (Scoring Rubric)
    add_heading2(doc, f"เกณฑ์การให้คะแนน (Scoring Rubric) {part_label} ตัวชี้วัดที่ {indicator_no}")
    rubric_rows = []
    for score in range(1, num_criteria + 1):
        rubric_rows.append([f"{score} คะแนน", f"[เกณฑ์การพิจารณาผลการปฏิบัติงาน/ผลลัพธ์ระดับคะแนน {score} "
                                                f"— คัดลอกข้อความจากแบบประเมิน PA 4 ของ ก.ค.ศ. ตามตำแหน่ง/วิทยฐานะจริง]"])
    add_table_grid(doc, headers=["คะแนน", "เกณฑ์การพิจารณาผลการปฏิบัติงาน / ผลลัพธ์"],
                    rows=rubric_rows, center_cols=[0])

    # 4) ตารางสรุปความสอดคล้อง
    add_heading2(doc, f"ตารางสรุป การวิเคราะห์ความสอดคล้องระหว่างเกณฑ์ PA 4 กับกลยุทธ์ของผู้จัดทำ")
    corr_rows = []
    for i in range(1, num_criteria + 1):
        corr_rows.append([f"{i}) [ข้อความเกณฑ์การพิจารณาข้อที่ {i} ตามแบบประเมิน]",
                           f"[การดำเนินงานตามกลยุทธ์ที่สอดคล้องกับเกณฑ์ข้อที่ {i}]",
                           f"[ผลลัพธ์ที่เกิดขึ้น]",
                           f"[หลักฐานอ้างอิง]"])
    add_table_grid(doc, headers=["เกณฑ์การพิจารณาผลการปฏิบัติตาม", "การดำเนินงานตามกลยุทธ์",
                                  "ผลลัพธ์", "หลักฐานอ้างอิง"], rows=corr_rows)

    # 5) การดำเนินงานตามกลยุทธ์ (PDCA)
    add_heading2(doc, "การดำเนินงานตามกลยุทธ์ที่ 1  [ชื่อกลยุทธ์]")
    add_paragraph_thai(doc, "ข้าพเจ้า [คำนำหน้า ชื่อ-นามสกุล] ตำแหน่ง [ตำแหน่ง] วิทยฐานะ [วิทยฐานะ] "
                             "มีกระบวนการดำเนินงานเพื่อให้เกิดผลลัพธ์ตามเกณฑ์การประเมิน ดังนี้")
    for i in range(1, num_criteria + 1):
        add_paragraph_thai(doc, f"•  1.{i}  [รายละเอียดกิจกรรม/ผลการดำเนินงานย่อยข้อที่ {i}]")

    add_paragraph_thai(doc, "ข้าพเจ้ามีกระบวนการในการดำเนินงาน ดังนี้")
    add_heading3(doc, "1.  ขั้นเตรียมการ")
    add_paragraph_thai(doc, "[บรรยายการศึกษาสภาพปัญหา/ความต้องการจำเป็น และการออกแบบกลยุทธ์/แผนงานก่อนดำเนินการ "
                             "พร้อมอ้างอิงหลักฐาน]")
    add_heading3(doc, "2.  ขั้นดำเนินการ")
    add_paragraph_thai(doc, "[บรรยายภาพรวมการนำแผนสู่การปฏิบัติ ตามวงจรคุณภาพ PDCA ดังนี้]")
    for code, label, note in [
        ("2.1", "ขั้นวางแผน (Plan)", "[รายละเอียดขั้นวางแผน]"),
        ("2.2", "ขั้นปฏิบัติตามแผน (Do)", "[รายละเอียดขั้นปฏิบัติ]"),
        ("2.3", "ขั้นตรวจสอบ (Check)", "[รายละเอียดขั้นตรวจสอบ/ติดตาม/ประเมินผล]"),
        ("2.4", "ขั้นพัฒนาปรับปรุง (Act)", "[รายละเอียดการนำผลประเมินไปปรับปรุง]"),
    ]:
        add_paragraph_thai(doc, f"•  {code}  {label}  {note}")
    add_heading3(doc, "3.  ขั้นจัดทำรายงานประเมินผล")
    add_paragraph_thai(doc, "[บรรยายการจัดทำรายงานสรุปผลการดำเนินงาน และการนำเสนอ/เผยแพร่ผลงาน]")

    # 6) ตารางที่ 1
    add_heading3(doc, f"ตารางที่ 1  การวิเคราะห์ความสอดคล้องกับเกณฑ์การให้คะแนน ตัวชี้วัดที่ {indicator_no}")
    t1_rows = []
    for i in range(1, num_criteria + 1):
        t1_rows.append([str(i), f"[เกณฑ์การพิจารณาตาม PA 4 ข้อที่ {i}]",
                         f"[ผลการปฏิบัติงานและหลักฐานเชิงประจักษ์ข้อที่ {i}]"])
    add_table_grid(doc, headers=["ข้อ", "เกณฑ์การพิจารณาตาม PA 4", "ผลการปฏิบัติงานและหลักฐานเชิงประจักษ์"],
                    rows=t1_rows, center_cols=[0])

    # 7) ผลลัพธ์ + ตารางที่ 2
    add_heading2(doc, "ผลลัพธ์")
    add_paragraph_thai(doc, "[สรุปผลลัพธ์ภาพรวมที่เกิดขึ้นจากการดำเนินงานตามกลยุทธ์นี้ "
                             "พร้อมตัวเลข/ร้อยละสนับสนุน]")
    add_heading3(doc, "ตารางที่ 2  ผลการดำเนินงาน (ตัวอย่างรูปแบบตาราง — แก้ไขหัวข้อ/จำนวนแถวตามจริง)")
    add_disclaimer_note(doc)
    t2_rows = []
    for i in range(1, num_criteria + 1):
        t2_rows.append([str(i), f"[กิจกรรมที่ {i}]", "[กลุ่มเป้าหมาย]", "[จำนวนที่เข้าร่วม]",
                         "0.00", "[ผลการดำเนินงาน]"])
    add_table_grid(doc, headers=["ที่", "กิจกรรม", "กลุ่มเป้าหมาย", "เข้าร่วม", "ร้อยละ", "ผลการดำเนินงาน"],
                    rows=t2_rows, center_cols=[0, 4])
    add_paragraph_thai(doc, "จากตารางที่ 2 พบว่า [สรุปผลภาพรวมจากตาราง]")

    # 8) รายการหลักฐานเชิงประจักษ์รายข้อ
    for i in range(1, num_criteria + 1):
        add_heading3(doc, f"1.1.{i}  [ชื่อกิจกรรม/ผลลัพธ์ย่อยข้อที่ {i}]")
        add_paragraph_thai(doc, f"[บรรยายเชื่อมโยงว่ากิจกรรม/ผลลัพธ์ข้อที่ {i} สอดคล้องกับตัวชี้วัดที่ "
                                 f"{indicator_no} อย่างไร]")
        add_mixed_paragraph(doc, [("ผลการดำเนินงาน : ", True), (f"[ระบุผลเชิงประจักษ์/ตัวเลขยืนยันของข้อที่ {i}]", False)])
        add_photo_placeholder(doc)

    # 9) ตารางที่ 3
    add_heading3(doc, "ตารางที่ 3  สรุปผลการดำเนินงานเทียบค่าเป้าหมาย")
    add_disclaimer_note(doc)
    t3_rows = []
    for i in range(1, num_criteria + 1):
        t3_rows.append([f"{i}) [ประเด็นการพิจารณาข้อที่ {i}]", "0.00", "[เป้าหมาย]", "0.00", "[บรรลุ/ไม่บรรลุ]"])
    add_table_grid(doc, headers=["ประเด็นการพิจารณา", "ข้อมูลฐาน (ร้อยละ)", "ค่าเป้าหมาย (ร้อยละ)",
                                  "ผลการประเมิน (ร้อยละ)", "สรุปผล"], rows=t3_rows, center_cols=[1, 2, 3, 4])
    add_paragraph_thai(doc, "จากตารางที่ 3 พบว่า [สรุปผลภาพรวมเทียบค่าเป้าหมาย]")
    add_paragraph_thai(doc, "ข้อสังเกตที่แสดงถึงระดับการปฏิบัติที่คาดหวังของวิทยฐานะที่ขอรับการประเมิน คือ "
                             "[ระบุข้อสังเกต/จุดเด่นที่ยกระดับจากวิทยฐานะเดิม]")

    # 10) ตารางแหล่งข้อมูลสำหรับตรวจสอบ (สำหรับผู้จัดทำเท่านั้น — ลบก่อนยื่นจริง)
    add_heading3(doc, "ตารางแหล่งข้อมูลสำหรับตรวจสอบ / แทนที่ตัวเลขสมมุติ (สำหรับผู้จัดทำ — ลบตารางนี้ออกก่อนยื่นเอกสารจริง)")
    add_paragraph_thai(doc, "ตารางนี้เป็นแนวทางบอกว่าแต่ละช่องตัวเลขในตารางที่ 2 และตารางที่ 3 ควรดึงข้อมูล"
                             "จากเอกสารหรือระบบใด เพื่อให้ตัวเลขที่กรอกตรวจสอบย้อนกลับได้")
    src_rows = []
    for i in range(min(num_source_rows, len(THAI_LETTERS))):
        src_rows.append([THAI_LETTERS[i], "[ชื่อเอกสาร/ระบบที่เป็นแหล่งข้อมูล]", "[วิธีนับ/หมายเหตุ]"])
    add_table_grid(doc, headers=["รหัส", "แหล่งข้อมูล", "วิธีนับ / หมายเหตุ"], rows=src_rows, center_cols=[0])

    # 11) ลงชื่อผู้รายงาน
    add_signature_block(doc)


# ---------------------------------------------------------------------------
# main
# ---------------------------------------------------------------------------

def main():
    doc = Document()
    set_document_defaults(doc)

    # ----- ปก (ไม่แสดงเลขหน้า) -----
    set_page_setup(doc.sections[0])
    build_cover_page(doc)

    # ----- ส่วนหน้า: เลขหน้าพยัญชนะไทย (ก ข ค ...) -----
    front = new_section(doc, WD_SECTION.NEW_PAGE)
    set_page_number_format(front, fmt="thaiLetters", start=1)
    add_page_number_footer(front)
    doc.add_page_break()
    add_toc_field(doc, caption="สารบัญ")

    # ----- เนื้อหา: เลขหน้าอารบิก (1 2 3 ...) -----
    body = new_section(doc, WD_SECTION.NEW_PAGE)
    set_page_number_format(body, fmt="decimal", start=1)
    add_page_number_footer(body)

    doc.add_page_break()
    build_part1_general_info(doc)

    build_indicator_block(
        doc,
        part_label="ส่วนที่ 2",
        domain_title="ผลการดำเนินงาน ด้านที่ 1  [ชื่อด้านที่ 1 ตามแบบประเมินตำแหน่ง/วิทยฐานะของผู้จัดทำ]",
        indicator_no=1,
        indicator_title="[ชื่อตัวชี้วัดที่ 1]",
    )

    build_indicator_block(
        doc,
        part_label="ส่วนที่ 3",
        domain_title="ผลการดำเนินงาน ด้านที่ 2  [ชื่อด้านที่ 2 ตามแบบประเมินตำแหน่ง/วิทยฐานะของผู้จัดทำ]",
        indicator_no=1,
        indicator_title="[ชื่อตัวชี้วัดที่ 1]",
    )

    doc.add_page_break()
    _center_run(doc, "ภาคผนวก", size=SIZE_H1, bold=True, space_after=18)
    add_paragraph_thai(doc, "(แนบหลักฐานประกอบ เช่น คำสั่งแต่งตั้ง ภาพถ่ายกิจกรรม รายงานการประชุม "
                             "แบบประเมิน/แบบสำรวจ และเอกสารอ้างอิงอื่น ๆ ที่กล่าวถึงในเล่ม)",
                        indent_first=1.25)

    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)
    print(f"บันทึกไฟล์แล้ว: {os.path.abspath(OUT_PATH)}")


if __name__ == "__main__":
    main()
