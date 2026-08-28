# -*- coding: utf-8 -*-
"""
build_5chapter_template.py
---------------------------
สร้างเทมเพลต Word (.docx) สำหรับ "รายงานวิจัย 5 บท" ตามมาตรฐานวิชาการไทย
(ปก / ใบรับรอง / บทคัดย่อ ไทย-อังกฤษ / กิตติกรรมประกาศ / สารบัญ / สารบัญตาราง /
สารบัญภาพ / บทที่ 1-5 / บรรณานุกรม / ภาคผนวก / ประวัติผู้วิจัย)

รัน:  python build_5chapter_template.py
ผลลัพธ์: ../output/5-Chapter-Research-Template.docx
"""
import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

from docx_common import (
    set_document_defaults, set_page_setup, set_page_number_format,
    add_page_number_footer, add_heading2, add_heading3, add_paragraph_thai,
    add_table_grid, add_toc_field, new_section, SIZE_H1, SIZE_BODY,
    _set_run_thai_font,
)

OUT_PATH = os.path.join(os.path.dirname(__file__), "..", "output",
                         "5-Chapter-Research-Template.docx")


def _center_run(doc, text, size=SIZE_BODY, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    _set_run_thai_font(r, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def build_cover_page(doc):
    for _ in range(3):
        doc.add_paragraph()
    _center_run(doc, "[ชื่อเรื่องงานวิจัยภาษาไทย]", size=22, bold=True, space_after=4)
    _center_run(doc, "[Research Title in English]", size=18, bold=True, space_after=40)
    _center_run(doc, "โดย", size=SIZE_H1, space_after=20)
    _center_run(doc, "[ชื่อ-นามสกุลผู้วิจัย]", size=SIZE_H1, bold=True, space_after=60)
    _center_run(doc, "รายงานการวิจัยฉบับนี้เป็นส่วนหนึ่งของการศึกษา", size=SIZE_BODY)
    _center_run(doc, "ตามหลักสูตร [ชื่อหลักสูตร/สาขาวิชา]", size=SIZE_BODY)
    _center_run(doc, "[ชื่อหน่วยงาน/สถานศึกษา/สังกัด]", size=SIZE_BODY)
    _center_run(doc, "ปีการศึกษา [พ.ศ. ....]", size=SIZE_BODY, space_after=0)


def build_approval_page(doc):
    _center_run(doc, "ใบรับรองรายงานการวิจัย", size=SIZE_H1, bold=True, space_after=24)
    add_paragraph_thai(doc, "เรื่อง [ชื่อเรื่องงานวิจัย] ฉบับนี้ ได้รับการตรวจสอบและอนุมัติให้เป็นส่วนหนึ่ง"
                             "ของการศึกษาตามหลักสูตร [ชื่อหลักสูตร] โดยคณะกรรมการสอบดังรายนามต่อไปนี้",
                        align=WD_ALIGN_PARAGRAPH.THAI_JUSTIFY, indent_first=1.25)
    doc.add_paragraph()
    for role in ["ประธานกรรมการ", "กรรมการ", "กรรมการและเลขานุการ"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run("......................................................  " + role)
        _set_run_thai_font(r, size=SIZE_BODY)
        p2 = doc.add_paragraph()
        r2 = p2.add_run("( ........................................................ )")
        _set_run_thai_font(r2, size=SIZE_BODY)
        p2.paragraph_format.left_indent = Cm(1.5)
        doc.add_paragraph()


def build_abstract_th(doc):
    _center_run(doc, "บทคัดย่อ", size=SIZE_H1, bold=True, space_after=12)
    for label, value in [
        ("ชื่อเรื่อง", "[ชื่อเรื่องงานวิจัยภาษาไทย]"),
        ("ผู้วิจัย", "[ชื่อ-นามสกุลผู้วิจัย]"),
        ("ปีการศึกษา", "[พ.ศ. ....]"),
    ]:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label} ")
        _set_run_thai_font(r1, size=SIZE_BODY, bold=True)
        r2 = p.add_run(value)
        _set_run_thai_font(r2, size=SIZE_BODY)
    doc.add_paragraph()
    add_paragraph_thai(doc, "(พิมพ์เนื้อหาบทคัดย่อภาษาไทยในส่วนนี้ ความยาวไม่เกิน 1 หน้ากระดาษ "
                             "ประกอบด้วย วัตถุประสงค์การวิจัย วิธีดำเนินการวิจัยโดยสรุป กลุ่มตัวอย่าง/"
                             "เครื่องมือที่ใช้ และผลการวิจัยที่สำคัญโดยสรุป)",
                        indent_first=1.25)
    p = doc.add_paragraph()
    r1 = p.add_run("คำสำคัญ ")
    _set_run_thai_font(r1, size=SIZE_BODY, bold=True)
    r2 = p.add_run("[คำสำคัญที่ 1], [คำสำคัญที่ 2], [คำสำคัญที่ 3]")
    _set_run_thai_font(r2, size=SIZE_BODY)


def build_abstract_en(doc):
    _center_run(doc, "ABSTRACT", size=SIZE_H1, bold=True, space_after=12)
    for label, value in [
        ("Title", "[Research Title in English]"),
        ("Researcher", "[Researcher Name]"),
        ("Academic Year", "[B.E. ....]"),
    ]:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}: ")
        _set_run_thai_font(r1, size=SIZE_BODY, bold=True)
        r2 = p.add_run(value)
        _set_run_thai_font(r2, size=SIZE_BODY)
    doc.add_paragraph()
    add_paragraph_thai(doc, "(Type the English abstract here — objectives, methodology, "
                             "sample/instruments, and key findings, within one page.)",
                        align=WD_ALIGN_PARAGRAPH.LEFT, indent_first=1.25)
    p = doc.add_paragraph()
    r1 = p.add_run("Keywords: ")
    _set_run_thai_font(r1, size=SIZE_BODY, bold=True)
    r2 = p.add_run("[keyword 1], [keyword 2], [keyword 3]")
    _set_run_thai_font(r2, size=SIZE_BODY)


def build_acknowledgement(doc):
    _center_run(doc, "กิตติกรรมประกาศ", size=SIZE_H1, bold=True, space_after=12)
    add_paragraph_thai(doc, "(พิมพ์ข้อความกิตติกรรมประกาศ ขอบคุณผู้มีส่วนช่วยเหลือให้งานวิจัยฉบับนี้"
                             "สำเร็จลุล่วง เช่น อาจารย์ที่ปรึกษา ผู้เชี่ยวชาญ กลุ่มตัวอย่าง หน่วยงานต้นสังกัด "
                             "และครอบครัว)", indent_first=1.25)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("[ชื่อ-นามสกุลผู้วิจัย]")
    _set_run_thai_font(r, size=SIZE_BODY)


def build_list_of_tables(doc):
    add_toc_field(doc, caption="สารบัญตาราง")
    add_paragraph_thai(doc, "(หลังพิมพ์เนื้อหาและใส่คำอธิบายภาพ/ตาราง (Caption) ครบแล้ว "
                             "ให้เลือกเมนู References > Insert Table of Figures โดยเลือก Caption label "
                             "เป็น \"ตาราง\" เพื่อสร้างสารบัญตารางอัตโนมัติ)",
                        align=WD_ALIGN_PARAGRAPH.LEFT, size=13)


def build_list_of_figures(doc):
    add_toc_field(doc, caption="สารบัญภาพ")
    add_paragraph_thai(doc, "(ทำเช่นเดียวกับสารบัญตาราง แต่เลือก Caption label เป็น \"ภาพ\")",
                        align=WD_ALIGN_PARAGRAPH.LEFT, size=13)


# ---------------------------------------------------------------------------
# บทที่ 1-5
# ---------------------------------------------------------------------------

def _chapter_heading(doc, chapter_no_th, title, first=False):
    if not first:
        doc.add_page_break()
    _center_run(doc, f"บทที่ {chapter_no_th}", size=SIZE_H1, bold=True, space_after=2)
    _center_run(doc, title, size=SIZE_H1, bold=True, space_after=18)


def _placeholder(doc, instruction):
    add_paragraph_thai(doc, f"({instruction})", indent_first=1.25)


def build_chapter1(doc):
    _chapter_heading(doc, "1", "บทนำ", first=True)
    for head, note in [
        ("1.1 ความเป็นมาและความสำคัญของปัญหา",
         "บรรยายที่มา สภาพปัญหา และเหตุผลที่ต้องทำวิจัยเรื่องนี้ โดยอ้างอิงข้อมูล/สถิติ/แนวคิดสนับสนุน"),
        ("1.2 วัตถุประสงค์ของการวิจัย",
         "ระบุวัตถุประสงค์เป็นข้อ ๆ ให้สอดคล้องกับชื่อเรื่องและคำถามวิจัย"),
        ("1.3 สมมติฐานของการวิจัย (ถ้ามี)",
         "ระบุสมมติฐานที่คาดการณ์ผลการวิจัย ในกรณีที่เป็นการวิจัยเชิงทดลอง/เชิงปริมาณ"),
        ("1.4 ขอบเขตของการวิจัย",
         "ระบุขอบเขตด้านเนื้อหา ประชากร/กลุ่มตัวอย่าง ตัวแปร และระยะเวลาที่ศึกษา"),
        ("1.5 นิยามศัพท์เฉพาะ",
         "ให้คำนิยามเชิงปฏิบัติการของตัวแปร/คำศัพท์สำคัญที่ใช้ในงานวิจัย"),
        ("1.6 ประโยชน์ที่คาดว่าจะได้รับ",
         "ระบุประโยชน์เชิงวิชาการและเชิงปฏิบัติที่จะเกิดขึ้นจากผลการวิจัย"),
    ]:
        add_heading2(doc, head)
        _placeholder(doc, note)


def build_chapter2(doc):
    _chapter_heading(doc, "2", "เอกสารและงานวิจัยที่เกี่ยวข้อง")
    for head, note in [
        ("2.1 แนวคิด ทฤษฎีที่เกี่ยวข้อง",
         "สังเคราะห์แนวคิด/ทฤษฎีหลักที่เป็นกรอบของงานวิจัย พร้อมอ้างอิงแหล่งที่มา"),
        ("2.2 งานวิจัยที่เกี่ยวข้อง",
         "สรุปงานวิจัยในและต่างประเทศที่เกี่ยวข้อง เรียงตามลำดับเวลาหรือประเด็น"),
        ("2.3 กรอบแนวคิดในการวิจัย",
         "นำเสนอกรอบแนวคิด (Conceptual Framework) เชื่อมโยงตัวแปรต้น-ตาม อาจแทรกภาพประกอบ"),
    ]:
        add_heading2(doc, head)
        _placeholder(doc, note)

    add_heading3(doc, "ตารางที่ 2.1  สรุปงานวิจัยที่เกี่ยวข้อง (ตัวอย่างรูปแบบตาราง)")
    add_table_grid(
        doc,
        headers=["ผู้วิจัย/ปี", "วัตถุประสงค์", "วิธีดำเนินการ", "ผลการวิจัยที่สำคัญ"],
        rows=[
            ["[ผู้แต่ง, ปีที่พิมพ์]", "[วัตถุประสงค์โดยย่อ]", "[กลุ่มตัวอย่าง/เครื่องมือ]", "[ผลสรุปที่เกี่ยวข้อง]"],
            ["[ผู้แต่ง, ปีที่พิมพ์]", "[วัตถุประสงค์โดยย่อ]", "[กลุ่มตัวอย่าง/เครื่องมือ]", "[ผลสรุปที่เกี่ยวข้อง]"],
        ],
    )


def build_chapter3(doc):
    _chapter_heading(doc, "3", "วิธีดำเนินการวิจัย")
    for head, note in [
        ("3.1 ประชากรและกลุ่มตัวอย่าง",
         "ระบุประชากร วิธีสุ่มตัวอย่าง ขนาดกลุ่มตัวอย่าง และเกณฑ์การคัดเลือก"),
        ("3.2 เครื่องมือที่ใช้ในการวิจัย",
         "อธิบายลักษณะเครื่องมือ (แบบสอบถาม/แบบทดสอบ/แบบสัมภาษณ์) และคุณภาพเครื่องมือ "
         "(ความตรง/ความเที่ยง) พร้อมผลการตรวจสอบโดยผู้เชี่ยวชาญ"),
        ("3.3 การเก็บรวบรวมข้อมูล",
         "อธิบายขั้นตอนและระยะเวลาการเก็บข้อมูลภาคสนาม"),
        ("3.4 การวิเคราะห์ข้อมูล",
         "ระบุวิธีวิเคราะห์ข้อมูลเชิงปริมาณ/เชิงคุณภาพที่ใช้ตอบวัตถุประสงค์แต่ละข้อ"),
        ("3.5 สถิติที่ใช้ในการวิจัย",
         "ระบุสถิติบรรยาย/สถิติอ้างอิงที่ใช้ พร้อมเกณฑ์การแปลผล"),
    ]:
        add_heading2(doc, head)
        _placeholder(doc, note)


def build_chapter4(doc):
    _chapter_heading(doc, "4", "ผลการวิเคราะห์ข้อมูล")
    add_heading2(doc, "4.1 สัญลักษณ์ที่ใช้ในการนำเสนอผลการวิเคราะห์ข้อมูล")
    _placeholder(doc, "ระบุสัญลักษณ์ทางสถิติที่ใช้ เช่น x̄ แทนค่าเฉลี่ย, S.D. แทนส่วนเบี่ยงเบนมาตรฐาน")

    add_heading2(doc, "4.2 ผลการวิเคราะห์ข้อมูล")
    _placeholder(doc, "นำเสนอผลการวิเคราะห์ข้อมูลเรียงตามวัตถุประสงค์/คำถามวิจัยแต่ละข้อ "
                       "ประกอบตารางและการแปลผล")

    add_heading3(doc, "ตารางที่ 4.1  ตัวอย่างรูปแบบตารางแสดงผลการวิเคราะห์ข้อมูล")
    add_table_grid(
        doc,
        headers=["รายการ", "จำนวน (n)", "ค่าเฉลี่ย (x̄)", "S.D.", "ระดับ"],
        rows=[
            ["[รายการที่ 1]", "0", "0.00", "0.00", "[ระดับ]"],
            ["[รายการที่ 2]", "0", "0.00", "0.00", "[ระดับ]"],
            ["รวม", "0", "0.00", "0.00", "[ระดับ]"],
        ],
        center_cols=[1, 2, 3, 4],
    )


def build_chapter5(doc):
    _chapter_heading(doc, "5", "สรุปผล อภิปรายผล และข้อเสนอแนะ")
    for head, note in [
        ("5.1 สรุปผลการวิจัย",
         "สรุปผลการวิจัยตามลำดับวัตถุประสงค์/คำถามวิจัยโดยย่อ ไม่ตีความเพิ่มเติม"),
        ("5.2 อภิปรายผล",
         "อภิปรายผลการวิจัยโดยเชื่อมโยงกับแนวคิด ทฤษฎี และงานวิจัยที่เกี่ยวข้องในบทที่ 2 "
         "พร้อมให้เหตุผลสนับสนุนหรือขัดแย้ง"),
        ("5.3 ข้อเสนอแนะ",
         "แบ่งเป็นข้อเสนอแนะเชิงนโยบาย/การนำผลไปใช้ และข้อเสนอแนะสำหรับการวิจัยครั้งต่อไป"),
    ]:
        add_heading2(doc, head)
        _placeholder(doc, note)


def build_bibliography(doc):
    doc.add_page_break()
    _center_run(doc, "บรรณานุกรม", size=SIZE_H1, bold=True, space_after=18)
    _placeholder(doc, "เรียงรายการอ้างอิงตามลำดับตัวอักษร ตามรูปแบบการอ้างอิงที่สถาบันกำหนด "
                       "(เช่น APA 7th) โดยแยกภาษาไทยและภาษาอังกฤษ")


def build_appendix(doc):
    doc.add_page_break()
    _center_run(doc, "ภาคผนวก", size=SIZE_H1, bold=True, space_after=18)
    for label in ["ภาคผนวก ก  เครื่องมือที่ใช้ในการวิจัย",
                  "ภาคผนวก ข  รายนามผู้เชี่ยวชาญตรวจสอบเครื่องมือ",
                  "ภาคผนวก ค  หนังสือขออนุญาต/หนังสือรับรองจริยธรรมการวิจัย"]:
        add_heading2(doc, label)
        _placeholder(doc, "แนบเอกสาร/รูปภาพที่เกี่ยวข้องในส่วนนี้")


def build_researcher_cv(doc):
    doc.add_page_break()
    _center_run(doc, "ประวัติผู้วิจัย", size=SIZE_H1, bold=True, space_after=18)
    for label, value in [
        ("ชื่อ-นามสกุล", "[ชื่อ-นามสกุล]"),
        ("วัน เดือน ปีเกิด", "[วัน เดือน ปี]"),
        ("ประวัติการศึกษา", "[วุฒิการศึกษา, สถาบัน, ปีที่สำเร็จ]"),
        ("ตำแหน่งปัจจุบัน", "[ตำแหน่ง/สถานที่ทำงาน]"),
        ("ผลงานทางวิชาการ (ถ้ามี)", "[ระบุผลงาน]"),
    ]:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label} ")
        _set_run_thai_font(r1, size=SIZE_BODY, bold=True)
        r2 = p.add_run(value)
        _set_run_thai_font(r2, size=SIZE_BODY)


def main():
    doc = Document()
    set_document_defaults(doc)

    # ----- ส่วนที่ 0: ปก (ไม่แสดงเลขหน้า) -----
    set_page_setup(doc.sections[0])
    build_cover_page(doc)

    # ----- ส่วนที่ 1: ส่วนหน้า เลขหน้าพยัญชนะไทย (ก ข ค ...) -----
    front = new_section(doc, WD_SECTION.NEW_PAGE)
    set_page_number_format(front, fmt="thaiLetters", start=1)
    add_page_number_footer(front)

    doc.add_page_break()
    build_approval_page(doc)
    doc.add_page_break()
    build_abstract_th(doc)
    doc.add_page_break()
    build_abstract_en(doc)
    doc.add_page_break()
    build_acknowledgement(doc)
    doc.add_page_break()
    add_toc_field(doc, caption="สารบัญ")
    doc.add_page_break()
    build_list_of_tables(doc)
    doc.add_page_break()
    build_list_of_figures(doc)

    # ----- ส่วนที่ 2: เนื้อหา เลขหน้าอารบิก (1 2 3 ...) -----
    body = new_section(doc, WD_SECTION.NEW_PAGE)
    set_page_number_format(body, fmt="decimal", start=1)
    add_page_number_footer(body)

    build_chapter1(doc)
    build_chapter2(doc)
    build_chapter3(doc)
    build_chapter4(doc)
    build_chapter5(doc)
    build_bibliography(doc)
    build_appendix(doc)
    build_researcher_cv(doc)

    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)
    print(f"บันทึกไฟล์แล้ว: {os.path.abspath(OUT_PATH)}")


if __name__ == "__main__":
    main()
